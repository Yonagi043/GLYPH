"""Build the GLYPH v1.2.0 group-meeting briefing DOCX."""
from __future__ import annotations

import csv
import json
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
RUN = ROOT / "data/processed/visual_features_v1/runs/render_551362ca0ff22f33"
OUT = ROOT / "docs/GLYPH_v1.2.0_group_meeting_brief.docx"
ASSET_DIR = ROOT / "docs/meeting_brief_assets"
REPO_URL = "https://github.com/Yonagi043/GLYPH"

NAVY = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
GRAY = RGBColor(85, 85, 85)
LIGHT = "F2F4F7"
PALE_BLUE = "E8EEF5"
WHITE = RGBColor(255, 255, 255)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    if isinstance(fill, RGBColor):
        fill = str(fill)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=130, bottom=90, end=130):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v)); node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths_dxa):
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW"); tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa))); tbl_w.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol"); col.set(qn("w:w"), str(width)); grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW"); tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width)); tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_font(run, name="Calibri", size=11, color=None, bold=None, italic=None):
    # LibreOffice's headless profile does not reliably honor eastAsia-only
    # font assignments, so Chinese runs explicitly use a Unicode-capable face
    # for all font slots.
    if any(ord(ch) > 127 for ch in run.text):
        name = "Arial Unicode MS"
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    rfonts.set(qn("w:ascii"), name); rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), "Arial Unicode MS"); rfonts.set(qn("w:cs"), "Arial Unicode MS"); rfonts.set(qn("w:hint"), "eastAsia")
    lang = rpr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang"); rpr.append(lang)
    lang.set(qn("w:val"), "en-US"); lang.set(qn("w:eastAsia"), "zh-CN")
    run.font.size = Pt(size)
    if color is not None: run.font.color.rgb = color
    if bold is not None: run.bold = bold
    if italic is not None: run.italic = italic


def style_document(doc: Document):
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Inches(8.5), Inches(11)
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
    sec.header_distance = sec.footer_distance = Inches(0.492)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"; normal.font.size = Pt(11); normal.font.color.rgb = RGBColor(30, 30, 30)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri"); normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri"); normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in (("Heading 1",16,BLUE,16,8),("Heading 2",13,BLUE,12,6),("Heading 3",12,RGBColor(31,77,120),8,4)):
        st=styles[name]; st.font.name="Calibri"; st.font.size=Pt(size); st.font.bold=True; st.font.color.rgb=color
        st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri"); st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri"); st._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
        st.paragraph_format.space_before=Pt(before); st.paragraph_format.space_after=Pt(after); st.paragraph_format.keep_with_next=True
    header = sec.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header.add_run("GLYPH  |  VISUAL FEATURES V1.2.0")
    set_font(r, size=8.5, color=GRAY, bold=True)
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("Internal group meeting brief  |  2026-08-30")
    set_font(r, size=8.5, color=GRAY)


def add_title(doc):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(4)
    r = p.add_run("GLYPH")
    set_font(r, size=27, color=NAVY, bold=True)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(3)
    r = p.add_run("第一版视觉特征数据基础")
    set_font(r, size=20, color=RGBColor(35, 35, 35), bold=True)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(16)
    r = p.add_run("组会简报 · protocol visual_features_v1.2.0")
    set_font(r, size=12, color=GRAY)
    table = doc.add_table(rows=4, cols=2); table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_widths(table, [1800, 7560])
    metadata = [("仓库", REPO_URL), ("参考运行", "render_551362ca0ff22f33"), ("数据规模", "140 个刺激 · 280 条视觉特征记录"), ("当前状态", "技术链完成；等待两名组员独立完成 fixture 审查")]
    for row, (label, value) in zip(table.rows, metadata):
        set_cell_shading(row.cells[0], PALE_BLUE)
        for idx, text in enumerate((label, value)):
            row.cells[idx].text = ""
            p = row.cells[idx].paragraphs[0]
            r = p.add_run(text); set_font(r, size=10.5, color=NAVY if idx == 0 else RGBColor(30,30,30), bold=(idx==0))
    doc.add_paragraph()
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(8)
    r = p.add_run("一句话结论：")
    set_font(r, size=12, color=BLUE, bold=True)
    r = p.add_run("我们已经把“字体审美”第一步收紧成一套可复现、可审查、可继续扩展的视觉测量基础；它不是审美真值，也没有替代真人实验。")
    set_font(r, size=12, color=RGBColor(30,30,30))


def add_metrics(doc):
    doc.add_heading("1. 我们现在拿到了什么", level=1)
    table = doc.add_table(rows=2, cols=4); table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_widths(table, [2340,2340,2340,2340])
    headers = ["140", "280", "7", "0"]
    labels = ["唯一刺激", "特征记录", "开放字体资产", "主运行失败"]
    for i, (value, label) in enumerate(zip(headers, labels)):
        cell = table.rows[0].cells[i]; set_cell_shading(cell, NAVY); cell.text = ""
        p=cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(value); set_font(r,size=22,color=WHITE,bold=True)
        cell = table.rows[1].cells[i]; set_cell_shading(cell, LIGHT); cell.text = ""
        p=cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(label); set_font(r,size=10.5,color=NAVY,bold=True)
    doc.add_paragraph("数据和运行产物都带有 manifest、字体哈希、许可证状态、PNG/mask 哈希、QC 报告和 checksums；旧版本运行已移入 history，不与当前协议混用。")
    doc.add_heading("2. 研究矩阵", level=1)
    table = doc.add_table(rows=1, cols=4); set_table_widths(table,[1800,3000,2760,1800])
    for cell, text in zip(table.rows[0].cells,["文字系统","内容","基线字体","WP4 汉字范例"]):
        set_cell_shading(cell, PALE_BLUE); cell.text=""; r=cell.paragraphs[0].add_run(text); set_font(r,size=10,color=NAVY,bold=True)
    matrix=[("Latn","A E H M O R S T + 2 序列","Noto Sans","—"),("Hani","永 山 水 中 文 門 木 人 + 2 序列","Noto Sans CJK SC","Serif / Iansui / Marker Gothic"),("Kana","あ い す ね の ま る よ + 2 序列","Noto Sans CJK JP","—"),("Hang","가 나 다 라 마 바 사 아 + 2 序列","Noto Sans KR","—")]
    for item in matrix:
        cells=table.add_row().cells
        for cell,text in zip(cells,item): cell.text=""; r=cell.paragraphs[0].add_run(text); set_font(r,size=9.5,color=RGBColor(30,30,30))
    doc.add_paragraph("每个条件使用两个 profile：bbox_height_matched（320 px）和 ink_area_matched（0.050）；所有字体与内容均固定，不根据结果替换。")


def add_pipeline(doc):
    doc.add_page_break()
    doc.add_heading("3. 标准化管线", level=1)
    steps=[("资产登记","检查许可证、Unicode 覆盖、变量轴、彩色/位图表和 SHA-256。"),("整形与渲染","NFC + HarfBuzz cluster 校验；Pillow/FreeType BASIC 固定栅格化。"),("归一化","只做等比例缩放；固定画布、锚点、黑字白底和二值阈值 128。"),("测量","灰度和二值各一份；输出密度、比例、几何、笔势、排版、重心、节奏、统一度代理量。"),("QC 与发布","重复哈希、敏感性、许可证和人工 fixture 审查；只有全部通过才能 release。")]
    table=doc.add_table(rows=1, cols=2); set_table_widths(table,[1900,7460])
    for c,t in zip(table.rows[0].cells,["阶段","固定动作"]): set_cell_shading(c,NAVY); c.text=""; r=c.paragraphs[0].add_run(t); set_font(r,size=10.5,color=WHITE,bold=True)
    for n,(stage,desc) in enumerate(steps,1):
        cells=table.add_row().cells; cells[0].text=""; r=cells[0].paragraphs[0].add_run(f"{n}. {stage}"); set_font(r,size=10.5,color=NAVY,bold=True); cells[1].text=""; r=cells[1].paragraphs[0].add_run(desc); set_font(r,size=10.5,color=RGBColor(30,30,30))
    doc.add_heading("4. 真实样本：同一汉字的四种汉字字体", level=1)
    contact = make_contact_sheet()
    doc.add_picture(str(contact), width=Inches(6.35))
    p=doc.paragraphs[-1]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    cap=doc.add_paragraph("图：同一字符“永”在四个汉字字体范例中的规范化输出；图像来自当前参考运行。")
    cap.alignment=WD_ALIGN_PARAGRAPH.CENTER; set_font(cap.runs[0],size=9,color=GRAY,italic=True)
    doc.add_paragraph("这一步得到的是可测的视觉差异：例如 ink coverage、bbox aspect ratio、对称性、闭合空间和骨架直/曲比例。它们是视觉代理量，不直接等于“更美”。")


def make_contact_sheet() -> Path:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    path = ASSET_DIR / "han_yong_font_comparison.png"
    items=[("Noto Sans CJK SC","stim_046feb329e243591.gray.png"),("Noto Serif SC","stim_43e4de4c0226c52b.gray.png"),("Bpmf Iansui","stim_7b6b8cc8ec5488f4.gray.png"),("LXGW Marker Gothic","stim_32f7caf9902568e8.gray.png")]
    sheet=Image.new("RGB",(1200,760),(248,249,250)); draw=ImageDraw.Draw(sheet)
    for i,(label,file) in enumerate(items):
        x=(i%2)*600; y=(i//2)*380
        img=Image.open(RUN/"rendered"/file).convert("L").resize((560,280))
        sheet.paste(Image.merge("RGB",(img,img,img)),(x+20,y+48))
        draw.text((x+24,y+16),label,fill=(11,37,69))
    sheet.save(path)
    return path


def add_review(doc):
    doc.add_page_break()
    doc.add_heading("5. 现在请组员做什么", level=1)
    p=doc.add_paragraph(); r=p.add_run("审查不是给字体打审美分，而是确认样本和协议没有明显的输入错误。"); set_font(r,size=12,color=NAVY,bold=True)
    doc.add_paragraph("请两名组员独立检查同一份 28 条 fixture 清单，并分别填写审查表。两人的记录不能先互相讨论后合并。")
    table=doc.add_table(rows=1, cols=3); set_table_widths(table,[2400,3480,3480])
    for c,t in zip(table.rows[0].cells,["检查项","要看什么","填写位置"]): set_cell_shading(c,PALE_BLUE); c.text=""; r=c.paragraphs[0].add_run(t); set_font(r,size=10.5,color=NAVY,bold=True)
    rows=[("刺激网格","28 条 fixture 是否齐全、内容和 profile 是否对应","visual_grid_pass"),("字符边界","字符顺序、边界和地区字形是否异常","character_boundary_pass"),("书体归类","字体角色/风格标签是否与图像一致","style_classification_pass")]
    for item in rows:
        cells=table.add_row().cells
        for cell,text in zip(cells,item): cell.text=""; r=cell.paragraphs[0].add_run(text); set_font(r,size=10.2,color=RGBColor(30,30,30))
    doc.add_heading("6. 审查完成后的动作", level=1)
    doc.add_paragraph("审查表位于 data/processed/visual_features_v1/human_review/fixture_review_records.csv。两名审查者都对 28 条记录填满且全部通过后，运行以下命令生成公开 release：")
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(0.25); r=p.add_run("uv run --locked --extra dev python -m glyph_features.cli release --run-id render_551362ca0ff22f33"); set_font(r,name="Courier New",size=10,color=NAVY)
    doc.add_heading("7. 边界与下一步", level=1)
    doc.add_paragraph("当前成果可以支撑下一轮的真人可读性/审美实验、文化叙事证据采集和跨字体配对比较，但不能直接声称某种字体“更美”。下一轮应继续沿用 stimulus_id、schema、许可证分层和 run 哈希，不重新发明字段。")
    p=doc.add_paragraph(); r=p.add_run("仓库地址："); set_font(r,size=12,color=BLUE,bold=True); r=p.add_run(REPO_URL); set_font(r,size=12,color=NAVY,bold=True)


def main():
    doc=Document(); style_document(doc); add_title(doc); add_metrics(doc); add_pipeline(doc); add_review(doc)
    OUT.parent.mkdir(parents=True, exist_ok=True); doc.save(OUT); print(OUT)


if __name__ == "__main__": main()
